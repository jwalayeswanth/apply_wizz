from __future__ import annotations

from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document


def _extract_text_from_docx(docx_path: Path) -> str:
    doc = Document(str(docx_path))

    chunks: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)

    # Best-effort: include table cell text (some resumes keep bullets/tables there).
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_texts.append(ct)
            if row_texts:
                chunks.append(" | ".join(row_texts))

    return "\n".join(chunks).strip()


def load_base_resume_text(data_dir: Path) -> str:
    """
    Load and return the candidate resume as plain text for the LLM prompt.
    """

    direct_path = data_dir / "candidate_resume.docx"
    if direct_path.exists():
        return _extract_text_from_docx(direct_path)

    # Fallback: look at Excel Resume Path column.
    excel_path = data_dir / "option2_job_links.xlsx"
    if not excel_path.exists():
        raise FileNotFoundError(f"Missing both candidate_resume.docx and Excel file: {excel_path}")

    df = pd.read_excel(excel_path)
    if "Resume Path" not in df.columns:
        raise ValueError("Excel is missing 'Resume Path' column required to locate resume docx.")

    resume_rel: Optional[str] = None
    for v in df["Resume Path"].tolist():
        if isinstance(v, str) and v.strip():
            resume_rel = v.strip()
            break
    if not resume_rel:
        raise ValueError("Could not find a non-empty value in Excel 'Resume Path' column.")

    resolved = (data_dir / resume_rel).resolve()
    if not resolved.exists():
        raise FileNotFoundError(f"Resume path from Excel does not exist: {resolved}")

    return _extract_text_from_docx(resolved)

