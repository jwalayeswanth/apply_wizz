from __future__ import annotations

from pathlib import Path

from docx import Document

from .types import JobDetails, TailoredResume
from .utils import ensure_dir


def generate_tailored_docx(
    job: JobDetails,
    tailored: TailoredResume,
    out_path: Path,
) -> Path:
    ensure_dir(out_path.parent)

    doc = Document()

    # Simple, consistent template.
    title = f"Tailored Resume: {job.title} ({job.company})"
    doc.add_paragraph(title)

    doc.add_paragraph("")  # spacer

    doc.add_paragraph("Summary").runs[0].bold = True
    doc.add_paragraph(tailored.summary)

    doc.add_paragraph("")
    doc.add_paragraph("Skills").runs[0].bold = True
    for s in tailored.skills:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Experience").runs[0].bold = True
    for b in tailored.experience_bullets:
        doc.add_paragraph(b, style="List Bullet")

    if tailored.projects_bullets:
        doc.add_paragraph("")
        doc.add_paragraph("Projects").runs[0].bold = True
        for b in tailored.projects_bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.save(str(out_path))
    return out_path

